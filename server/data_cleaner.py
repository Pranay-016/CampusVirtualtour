import json
import re
from docx import Document

def clean_text(text):
    return re.sub(r'\s+', ' ', str(text)).strip()

def parse_docx(file_path, output_path):
    doc = Document(file_path)
    
    departments = []
    
    # Extract departments from paragraphs
    for p in doc.paragraphs:
        txt = clean_text(p.text)
        match = re.match(r'DEPT\s*:\s*(.*)', txt, re.IGNORECASE)
        if match:
            departments.append(match.group(1).strip())
            
    print(f"Found {len(departments)} departments and {len(doc.tables)} tables.")
    
    faculty_data = []
    id_counter = 1
    
    # Process tables (we map by index, assuming one table per department)
    for i in range(min(len(departments), len(doc.tables))):
        dept = departments[i]
        table = doc.tables[i]
        
        # Skip header row if it exists
        start_row = 1 if len(table.rows) > 0 and 'name' in clean_text(table.rows[0].cells[0].text).lower() else 0
            
        for row in table.rows[start_row:]:
            if len(row.cells) >= 3:
                name = clean_text(row.cells[0].text)
                designation = clean_text(row.cells[1].text)
                qualification = clean_text(row.cells[2].text)
                
                if not name or name.lower() == 'name':
                    continue
                    
                if not designation:
                    designation = "Not specified"
                if not qualification:
                    qualification = "Not specified"
                    
                document_text = f"{name} is a {designation} in the {dept} department with a {qualification} qualification."
                
                faculty_data.append({
                    "id": f"faculty_{id_counter}",
                    "document": document_text,
                    "metadata": {
                        "name": name,
                        "department": dept,
                        "designation": designation,
                        "qualification": qualification
                    }
                })
                id_counter += 1
                
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(faculty_data, f, indent=4)
        
    print(f"Extracted {len(faculty_data)} faculty members.")

if __name__ == "__main__":
    parse_docx("All_Faculty_Data.docx", "faculty_data.json")
